from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    h.font.name = 'Calibri'
    if level == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
    else:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(12)

def add_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_font(cell, text, bold=False, size=9, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    if color:
        run.font.color.rgb = color

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_font(cell, h, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
        add_shading(cell, '0F172A')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            set_cell_font(cell, str(val), size=9)
            if ri % 2 == 1:
                add_shading(cell, 'F1F5F9')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

# ═══════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AI Predictive Maintenance System')
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Technical Architecture & Design Report')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x47, 0x52, 0x64)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('July 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 1.0')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    ('1.', 'Introduction — Use Case'),
    ('2.', 'System Architecture'),
    ('3.', 'Technology Stack'),
    ('4.', 'Data Flow'),
    ('5.', 'ML Pipeline Details'),
    ('6.', 'Database Schema'),
    ('7.', 'Frontend Pages'),
    ('8.', 'Reporting & Export'),
    ('9.', 'Theming System'),
    ('10.', 'Key Observations & Design Debt'),
    ('11.', 'Recommendations'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {title}')
    run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════
doc.add_heading('1. Introduction — Use Case', level=1)

doc.add_paragraph(
    'Aircraft engines generate continuous streams of sensor data during operation — temperature, '
    'pressure, rotational speeds, and vibration measurements. Unplanned engine failures are '
    'catastrophic: they cause flight delays, unscheduled grounding, expensive emergency repairs, '
    'and, in worst cases, safety incidents. The central challenge is predicting when an engine '
    'will fail so that maintenance can be scheduled proactively rather than reactively.'
)

doc.add_paragraph(
    'This system addresses that exact use case. It ingests 15 sensor measurements and 2 operational '
    'settings from turbofan engines (using the NASA CMAPSS benchmark dataset), runs them through '
    'a trained XGBoost Regressor, and outputs Remaining Useful Life (RUL) in cycles. The RUL is '
    'classified into three health tiers:'
)

bullets = [
    'Healthy (>80 cycles) — Continue normal operation',
    'Warning (30–80 cycles) — Schedule maintenance inspection',
    'Critical (<30 cycles) — Stop operation immediately'
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    'Every prediction is accompanied by a SHAP explanation (waterfall + bar plots) showing exactly '
    'which sensors drove the predicted RUL, turning a black-box ML output into an auditable, '
    'interpretable decision-support tool.'
)

doc.add_paragraph(
    'The application also provides a complete maintenance workflow: engineers can log maintenance '
    'records, view fleet health dashboards, drill into individual engine histories, and export '
    'reports in CSV, Excel, and PDF formats.'
)

# ═══════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════
doc.add_heading('2. System Architecture', level=1)

doc.add_paragraph(
    'The application follows a layered architecture with strict separation of concerns. '
    'Each layer has a well-defined responsibility and communicates through clear interfaces.'
)

doc.add_heading('2.1 Architectural Layers', level=2)

code = """┌─────────────────────────────────────────────────────┐
│                    Streamlit UI                      │
│  app/home.py (auth)  │  app/pages/*.py (5 pages)   │
├─────────────────────────────────────────────────────┤
│              Application / Service Layer             │
│  application/auth/     │  application/prediction/    │
│  application/reports/  │  (charts, export, theme)   │
├─────────────────────────────────────────────────────┤
│                   ML Engine                          │
│  ai/predict.py    │  ai/preprocessing.py            │
│  ai/explain.py    │  ai/model/*.pkl (3 artifacts)   │
├─────────────────────────────────────────────────────┤
│                 Data / Persistence                   │
│  database/models.py  │  database/queries.py          │
│  SQLite (maintenance.db)                             │
└─────────────────────────────────────────────────────┘"""
add_code_block(doc, code)

doc.add_heading('2.2 Key Design Decisions', level=2)

decisions = [
    ('Streamlit as the single-page-app framework', 'No separate frontend/backend, minimal boilerplate, ideal for ML dashboards'),
    ('SQLite for persistence', 'Zero-config, file-based, sufficient for single-team/single-instance deployment'),
    ('Service layer decouples UI from logic', 'PredictionService, AuthService, and report functions are testable independently'),
    ('Session state carries context', 'st.session_state carries user identity, theme preference, and transient results across pages'),
]
make_table(doc, ['Decision', 'Rationale'], decisions, col_widths=[6, 12])

# ═══════════════════════════════════════════════════════════
# 3. TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════
doc.add_heading('3. Technology Stack', level=1)

stack = [
    ('UI Framework', 'Streamlit', '>=1.58', 'Multi-page dashboard, session management, interactive widgets'),
    ('ML Model', 'XGBoost Regressor', '>=2.0', 'Gradient-boosted decision trees for RUL prediction'),
    ('Explainability', 'SHAP (TreeExplainer)', '>=0.51', 'Feature attribution for every prediction'),
    ('Charts', 'Plotly Express', '>=6.8', 'Interactive KPI charts (pie, bar, line, scatter)'),
    ('PDF Export', 'ReportLab', '>=5.0', 'Single-prediction PDF reports'),
    ('Excel Export', 'OpenPyXL', '>=3.1', 'Multi-sheet Excel workbook export'),
    ('Auth', 'bcrypt', '>=5.0', 'Password hashing and verification'),
    ('Data', 'Pandas / NumPy', '>=2.3 / >=2.4', 'DataFrame manipulation, numerical operations'),
    ('Scaling', 'scikit-learn StandardScaler', '>=1.9', 'Feature normalization'),
    ('Serialization', 'joblib', '>=1.5', 'Model persistence (.pkl)'),
    ('Database', 'SQLite (via sqlite3)', 'Built-in', 'Local relational storage'),
]
make_table(doc, ['Component', 'Technology', 'Version', 'Purpose'], stack, col_widths=[3.5, 4.5, 3, 7])

# ═══════════════════════════════════════════════════════════
# 4. DATA FLOW
# ═══════════════════════════════════════════════════════════
doc.add_heading('4. Data Flow', level=1)

doc.add_heading('4.1 Prediction Pipeline', level=2)

doc.add_paragraph(
    'The prediction pipeline transforms raw sensor inputs into actionable RUL predictions with '
    'the following steps:'
)

code = """User Input (manual or CSV)
    │
    ▼
validate_single_input() / validate_csv()
    │
    ▼
preprocess_single() / preprocess_batch()
    ├── Fill missing features with 0
    ├── Drop 7 unused sensors (NASA CMAPSS feature selection)
    ├── For CSV: group-by engine, take last cycle, parse numeric, drop NaN
    └── Scale via StandardScaler (fitted on training data)
    │
    ▼
XGBoost Regressor (xgboost_rul_model.pkl)
    │
    ▼
Predicted RUL (float) → get_health_status()
    ├── > 80  → "Healthy"   (green)
    ├── 30-80 → "Warning"   (amber)
    └── < 30  → "Critical"  (red)
    │
    ▼
SHAP TreeExplainer → waterfall + bar plots
    │
    ▼
Save: insert_prediction() + insert_sensor_readings()
    │
    ▼
Display: metric cards, SHAP plots, PDF download button"""
add_code_block(doc, code)

doc.add_heading('4.2 Input Features', level=2)

doc.add_paragraph('The model uses 19 features derived from the NASA CMAPSS dataset:')

doc.add_paragraph('Operational Settings:', style='List Bullet')
doc.add_paragraph('op_setting_1, op_setting_2  (op_setting_3 is dropped as redundant)', style='List Bullet 2')

doc.add_paragraph('Sensors (15 active):', style='List Bullet')
doc.add_paragraph('sensor_2, 3, 4, 6, 7, 8, 9, 11, 12, 13, 14, 15, 17, 20, 21', style='List Bullet 2')

doc.add_paragraph('Dropped sensors: sensor_1, 5, 10, 16, 18, 19 (based on CMAPSS research)', style='List Bullet')

doc.add_paragraph(
    'CSV batch input format: 26-column NASA-format files (unit, cycle, 3 settings, 21 sensors) '
    '— auto-parses to the 19-feature schema via group-by-unit with last-cycle extraction.'
)

# ═══════════════════════════════════════════════════════════
# 5. ML PIPELINE
# ═══════════════════════════════════════════════════════════
doc.add_heading('5. ML Pipeline Details', level=1)

doc.add_heading('5.1 Model Artifacts', level=2)

artifacts = [
    ('xgboost_rul_model.pkl', '~2 MB', 'Trained XGBoost Regressor (100–300 trees, max_depth ~6)'),
    ('standard_scaler.pkl', '~2 KB', 'Fitted StandardScaler (mean/std for each of 19 features)'),
    ('feature_names.pkl', '~1 KB', 'List of 19 feature column names'),
]
make_table(doc, ['File', 'Size', 'Description'], artifacts, col_widths=[5, 2.5, 10.5])

doc.add_heading('5.2 Predictor Class (ai/predict.py)', level=2)
doc.add_paragraph('predict_single(sensor_data: dict) → float: Scales input, runs model.predict()')
doc.add_paragraph('predict_batch(df: pd.DataFrame) → (np.array, pd.DataFrame): Groups by engine, takes last cycle, scales, batch predicts')

doc.add_heading('5.3 Explainability (ai/explain.py)', level=2)
doc.add_paragraph(
    'The Explainability class wraps shap.TreeExplainer(model) — the optimal explainer for '
    'tree-based models (exact computation, no approximation). The explain_single() method '
    'returns SHAP values, base value, and feature names. The Prediction page renders both '
    'waterfall plots (shap.plots.waterfall) and bar plots (shap.plots.bar) via matplotlib.'
)

doc.add_heading('5.4 Health Classification Thresholds', level=2)

thresholds = [
    ('> 80', 'Healthy', '"Continue normal operation."', '#2ecc71'),
    ('30 – 80', 'Warning', '"Schedule maintenance inspection."', '#f39c12'),
    ('< 30', 'Critical', '"Stop operation immediately."', '#e74c3c'),
]
make_table(doc, ['RUL Range', 'Status', 'Recommended Action', 'Color'],
           [(r[0], r[1], r[2], '') for r in thresholds],
           col_widths=[3, 3, 8, 4])

# ═══════════════════════════════════════════════════════════
# 6. DATABASE SCHEMA
# ═══════════════════════════════════════════════════════════
doc.add_heading('6. Database Schema', level=1)

doc.add_paragraph('Four tables in SQLite (database/maintenance.db):')

doc.add_heading('6.1 users', level=2)
make_table(doc, ['Column', 'Type', 'Constraints'],
           [('id', 'INTEGER', 'PK AUTOINCREMENT'),
            ('full_name', 'TEXT', 'NOT NULL'),
            ('username', 'TEXT', 'UNIQUE NOT NULL'),
            ('email', 'TEXT', 'UNIQUE NOT NULL'),
            ('password', 'TEXT', 'NOT NULL (bcrypt hash)'),
            ('role', 'TEXT', "CHECK('Admin','Engineer')"),
            ('created_at', 'TEXT', 'DEFAULT CURRENT_TIMESTAMP')],
           col_widths=[4, 4, 10])

doc.add_heading('6.2 predictions', level=2)
make_table(doc, ['Column', 'Type', 'Constraints'],
           [('id', 'INTEGER', 'PK AUTOINCREMENT'),
            ('engine_id', 'INTEGER', ''),
            ('user_id', 'INTEGER', 'FK → users(id) ON DELETE CASCADE'),
            ('prediction_date', 'TEXT', 'NOT NULL (ISO datetime)'),
            ('predicted_rul', 'REAL', 'NOT NULL'),
            ('health_status', 'TEXT', 'NOT NULL')],
           col_widths=[4, 4, 10])

doc.add_heading('6.3 sensor_readings', level=2)
make_table(doc, ['Column', 'Type', 'Constraints'],
           [('id', 'INTEGER', 'PK AUTOINCREMENT'),
            ('prediction_id', 'INTEGER', 'FK → predictions(id) ON DELETE CASCADE'),
            ('op_setting_1, op_setting_2', 'REAL', ''),
            ('sensor_2 … sensor_21', 'REAL', '15 sensor columns')],
           col_widths=[4, 4, 10])

doc.add_heading('6.4 maintenance', level=2)
make_table(doc, ['Column', 'Type', 'Constraints'],
           [('id', 'INTEGER', 'PK AUTOINCREMENT'),
            ('engine_id', 'INTEGER', 'NOT NULL'),
            ('user_id', 'INTEGER', 'FK → users(id) ON DELETE CASCADE'),
            ('maintenance_date', 'TEXT', 'NOT NULL'),
            ('maintenance_type', 'TEXT', 'NOT NULL'),
            ('notes', 'TEXT', ''),
            ('status', 'TEXT', "DEFAULT 'Completed'"),
            ('cost', 'REAL', 'DEFAULT 0'),
            ('downtime_hours', 'REAL', 'DEFAULT 0')],
           col_widths=[4, 4, 10])

# ═══════════════════════════════════════════════════════════
# 7. FRONTEND PAGES
# ═══════════════════════════════════════════════════════════
doc.add_heading('7. Frontend Pages', level=1)

pages = [
    ('Login / Signup', 'app/home.py', '/', 'Tabbed login/signup, bcrypt auth, role-based entry'),
    ('Prediction', 'app/pages/Prediction.py', '/Prediction', 'Manual mode (19 inputs + SHAP viz), CSV batch upload, PDF export'),
    ('History', 'app/pages/History.py', '/History', 'Per-user prediction history, engine filter, styled dataframe, CSV download'),
    ('Maintenance', 'app/pages/Maintenance.py', '/Maintenance', 'Fleet health overview, priority queue, engine cards with RUL bars, KPI stats'),
    ('Maintenance Records', 'app/pages/Maintenance_Records.py', '/Maintenance_Records', 'Log new record, engine drill-down, RUL trend chart with maintenance events'),
    ('Reports', 'app/pages/Reports.py', '/Reports', '6-tab analytics: KPIs, critical engines, 5 Plotly charts, full Excel export'),
]
make_table(doc, ['Page', 'File', 'Route', 'Key Features'], pages, col_widths=[3.5, 5.5, 2.5, 6.5])

doc.add_heading('7.1 Common UI Pattern', level=2)
doc.add_paragraph('Every authenticated page follows this pattern:')
add_code_block(doc, "st.set_page_config(...)\ninject_global_css()     # Apply theme CSS\nrender_sidebar()        # Navigation + theme toggle + logout\nif 'user' not in st.session_state: st.stop()  # Auth guard")

# ═══════════════════════════════════════════════════════════
# 8. REPORTING & EXPORT
# ═══════════════════════════════════════════════════════════
doc.add_heading('8. Reporting & Export', level=1)

doc.add_heading('8.1 Report Queries (application/reports/reports.py)', level=2)
doc.add_paragraph('13 SQL-based queries powering the Reports page:')

queries = [
    ('get_dashboard_kpis()', 'dict', 'User, prediction, maintenance, critical engine counts'),
    ('get_prediction_summary()', 'dict', 'Avg RUL, min/max RUL, total predictions'),
    ('get_critical_engines()', 'DataFrame', 'Engines with RUL < 30 or critical status'),
    ('get_recent_predictions(n)', 'DataFrame', 'Last n predictions by date DESC'),
    ('get_maintenance_history()', 'DataFrame', 'All maintenance records with engineer names'),
    ('get_engineer_activity()', 'DataFrame', 'Per-engineer prediction and maintenance counts'),
    ('get_health_status_distribution()', 'DataFrame', 'Health status → count'),
    ('get_prediction_trend()', 'DataFrame', 'Predictions per day over time'),
]
make_table(doc, ['Function', 'Returns', 'Description'], queries, col_widths=[6, 3.5, 8.5])

doc.add_heading('8.2 Chart Factories (application/reports/charts.py)', level=2)
doc.add_paragraph('5 Plotly chart types, all applying theme.apply_theme() for dark/light mode:')
charts = [
    ('Health Status Distribution', 'Donut pie chart with custom color map'),
    ('Prediction Trend', 'Spline line chart (predictions per day)'),
    ('Maintenance Type Distribution', 'Bar chart'),
    ('Engineer Activity', 'Grouped bar chart (predictions vs maintenance)'),
    ('Critical Engines', 'Horizontal bar chart (RUL per engine, color-coded)'),
]
make_table(doc, ['Chart', 'Description'], charts, col_widths=[6, 12])

doc.add_heading('8.3 Export Formats (application/reports/export.py)', level=2)
exports = [
    ('CSV', 'export_to_csv(df)', 'Single DataFrame → CSV bytes'),
    ('Multi-sheet Excel', 'export_multiple_reports(dict)', 'Multiple DataFrames → one .xlsx'),
    ('PDF', 'export_prediction_pdf(dict)', 'Single prediction report via ReportLab'),
]
make_table(doc, ['Format', 'Function', 'Content'], exports, col_widths=[4, 6, 8])

# ═══════════════════════════════════════════════════════════
# 9. THEMING SYSTEM
# ═══════════════════════════════════════════════════════════
doc.add_heading('9. Theming System', level=1)

doc.add_paragraph(
    'The app has a custom light/dark toggle stored in st.session_state.theme. '
    'The theming operates across three layers:'
)

doc.add_heading('9.1 CSS Architecture (app/utils/styles.py)', level=2)
doc.add_paragraph(
    'Two competing CSS systems exist in the same file (a design debt):'
)
doc.add_paragraph(
    'CSS variable (lines 15–731): Sophisticated CSS custom-property system with 80+ variables '
    'for both light (:root) and dark (html[data-theme="dark"]) themes. Includes neon text, '
    'glow borders, scan-line overlay, gauge circles, animated gradient borders, and pulsing '
    'status dots. This code is dead — lines 974–975 comment out both inject_theme_js() and '
    'st.markdown(CSS, ...).',
    style='List Bullet'
)
doc.add_paragraph(
    'inject_global_css() function (lines 738–973): The live styling function. Renders simpler '
    'inline <style> blocks based on st.session_state.theme. Dark mode uses #0F172A background, '
    '#111827 sidebar, #1E293B cards, #2563EB buttons.',
    style='List Bullet'
)

doc.add_heading('9.2 Streamlit Native Theme (.streamlit/config.toml)', level=2)
add_code_block(doc, """[theme]
primaryColor = "#00e5ff"       # Cyan accent
backgroundColor = "#0F172A"    # Dark page background
secondaryBackgroundColor = "#1E293B"  # Card/sidebar background
textColor = "#F8FAFC"          # Light text on dark
font = "sans serif"

[server]
maxUploadSize = 50""")

doc.add_heading('9.3 Plotly Theme (application/reports/theme.py)', level=2)
doc.add_paragraph(
    'get_plotly_template() returns plotly_dark with transparent backgrounds in dark mode, '
    'plotly_white in light mode. apply_theme(fig) applies the template and sets subtle grid lines. '
    'Dual color maps (HEALTH_COLORS for light, DARK_HEALTH_COLORS with bright neon for dark) '
    'ensure charts are legible in both modes.'
)

doc.add_heading('9.4 Theme Toggle (app/utils/sidebar.py:42-46)', level=2)
add_code_block(doc, """theme_icon = "🌙" if st.session_state.theme == "light" else "☀️"
theme_label = "Dark Mode" if st.session_state.theme == "light" else "Light Mode"
if st.button(f"{theme_icon} {theme_label}", key="nav_theme", width='stretch'):
    st.session_state.theme = "dark" if st.session_state.theme == "light" else "light"
    st.rerun()""")

# ═══════════════════════════════════════════════════════════
# 10. DESIGN DEBT
# ═══════════════════════════════════════════════════════════
doc.add_heading('10. Key Observations & Design Debt', level=1)

debt = [
    ('Dead CSS code', 'The CSS variable (700+ lines with elaborate custom-property system) is entirely unused. The live inject_global_css() is a separate, simpler implementation. The dead code should be removed or activated.'),
    ('Duplicate imports in styles.py', 'Lines 733 and 735 both re-import "import streamlit as st" unnecessarily.'),
    ('Session state ordering fragility', 'inject_global_css() is called before render_sidebar(). It uses st.session_state.get("theme", "dark") as fallback, which works but is fragile.'),
    ('theme.py still defaults to "light"', 'Line 25 uses st.session_state.get("theme", "light") — should be "dark" to match the new default.'),
    ('PredictionService lazy-loads SHAP', 'self.explainer = None is set on init; SHAP TreeExplainer loads on first prediction. Good optimization but first prediction will be slower.'),
    ('No training pipeline', 'The XGBoost model is pre-trained. No retraining, hyperparameter tuning, or experiment tracking. MLflow would add immediate value here.'),
    ('No automated tests', 'The project has no test files despite a well-structured service layer.'),
]
make_table(doc, ['Issue', 'Description'], debt, col_widths=[5, 13])

# ═══════════════════════════════════════════════════════════
# 11. RECOMMENDATIONS
# ═══════════════════════════════════════════════════════════
doc.add_heading('11. Recommendations', level=1)

doc.add_heading('11.1 MLflow Integration', level=2)
make_table(doc, ['Phase', 'Action', 'Location'],
           [('Tracking', 'Wrap training with mlflow.start_run() logging params, metrics, model artifact', 'ai/train.py (new)'),
            ('Registry', 'Register models to MLflow Model Registry, promote to Staging/Production', 'ai/train.py'),
            ('Inference', 'Load production model from registry via mlflow.xgboost.load_model()', 'ai/predict.py'),
            ('UI', 'mlflow ui for experiment comparison, parameter importance, run history', 'CLI')],
           col_widths=[3.5, 9, 5.5])

doc.add_heading('11.2 RAG System', level=2)
make_table(doc, ['Phase', 'Action', 'Stack'],
           [('Ingestion', 'Chunk maintenance records, PDF reports, NASA docs → embed → store', 'sentence-transformers + ChromaDB'),
            ('Retrieval', 'User query → embedding → top-k similar chunks', 'ChromaDB similarity search'),
            ('Generation', 'Prompt = context + query → local LLM response', 'Ollama (llama3/mistral) or OpenAI'),
            ('UI', 'Add app/pages/AIAssistant.py with chat interface', 'Streamlit chat elements')],
           col_widths=[3.5, 9, 5.5])

# ── Save ──────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), 'AI_Predictive_Maintenance_Technical_Report.docx')
doc.save(output_path)
print(f'Report saved to: {output_path}')
